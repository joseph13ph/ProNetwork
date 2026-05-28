from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / 'INFORME DE PRUEBAS DE SOFTWARE.docx'
OUTPUT_PATH = ROOT / 'INFORME DE PRUEBAS DE SOFTWARE actualizado.docx'
IMG_DIR = ROOT / 'docs' / 'screenshots'


def set_paragraph(paragraph, text, bold=False, center=False):
    paragraph.text = text
    if paragraph.runs:
        paragraph.runs[0].bold = bold
        paragraph.runs[0].font.size = Pt(11 if bold else 10)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def delete_table(table):
    tbl = table._tbl
    tbl.getparent().remove(tbl)


def move_table_after(table, paragraph):
    paragraph._p.addnext(table._tbl)


def set_table_cell_text(cell, text, font_size=9, bold=False):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.bold = bold


def add_image_to_cell(cell, caption, image_name, width=2.35):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption + '\n')
    run.bold = True
    run.font.size = Pt(8)
    image_path = IMG_DIR / image_name
    if image_path.exists():
        paragraph.add_run().add_picture(str(image_path), width=Inches(width))


def fill_plan_table(table, rows):
    headers = ['ID de caso de prueba', 'Tipo de prueba', 'ID de paso', 'Descripción del paso', 'Fecha de la prueba', 'Resultado Esperado']
    for col, header in enumerate(headers):
        set_table_cell_text(table.rows[0].cells[col], header, font_size=9, bold=True)
    for row_index, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            set_table_cell_text(table.rows[row_index].cells[col], value, font_size=8)


def fill_result_table(table, test_id, result, image_name, observations, caption):
    headers = ['ID de Prueba', 'Resultado', 'Evidencia', 'Observaciones']
    for col, header in enumerate(headers):
        set_table_cell_text(table.rows[0].cells[col], header, font_size=9, bold=True)
    set_table_cell_text(table.rows[1].cells[0], test_id)
    set_table_cell_text(table.rows[1].cells[1], result)
    set_table_cell_text(table.rows[1].cells[3], observations)
    add_image_to_cell(table.rows[1].cells[2], caption, image_name)


def main():
    doc = Document(str(DOCX_PATH))

    # Update section 4 and 5 text blocks.
    section4_texts = [
        '4. PLAN DE PRUEBAS',
        'Basado en la interfaz real de ProConnect: Feed, Conexiones, Empleos, Mensajes, botón "Publicar ahora" y campo de texto "Comparte un logro, una vacante o una idea que pueda inspirar a otros...".',
        'Pruebas de Aceptación (Usuario) - 3 pruebas',
        'Objetivo: validar flujos completos desde la perspectiva del usuario final.',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        ''
    ]
    for idx, text in enumerate(section4_texts, start=109):
        set_paragraph(doc.paragraphs[idx], text, bold=(idx in (109, 112, 118) or idx == 109), center=False)

    section5_texts = [
        '5. EJECUCIÓN DE PRUEBAS',
        'La ejecución de pruebas de ProConnect se realizó en un entorno local para validar el comportamiento funcional de la interfaz real, la navegación entre módulos y las acciones de autenticación, publicación y cierre de sesión.',
        'Entorno de Ejecución de las Pruebas',
        'Frontend: React 18 + Vite, ejecutado en http://localhost:5173.',
        'Backend: Node.js + Express, ejecutado en http://localhost:5001.',
        'Base de datos: SQLite local en backend/data/proconnect.sqlite.',
        'Navegador: Google Chrome.',
        'Cobertura: login, feed, navegación entre módulos, publicación de contenido, validaciones y cierre de sesión.',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        ''
    ]
    for idx, text in enumerate(section5_texts, start=122):
        set_paragraph(doc.paragraphs[idx], text, bold=(idx == 122 or idx == 124), center=False)

    # Repurpose the four plan tables and move them into section 4.
    fill_plan_table(doc.tables[1], [
        ['PU-001', 'Usuario', 'ACC-01', 'Validar flujo completo: Inicio de sesión → Feed', '25/05/2026', 'El sistema redirige al Feed y muestra el mensaje "Bienvenido, [nombre]" con el campo de publicación.'],
        ['PU-002', 'Usuario', 'ACC-02', 'Validar flujo completo: Creación de publicación', '25/05/2026', 'Al hacer clic en "Publicar ahora", la publicación aparece en el feed con foto de perfil, nombre y texto.'],
        ['PU-003', 'Usuario', 'ACC-03', 'Validar flujo completo: Navegación entre módulos', '25/05/2026', 'Al hacer clic en "Empleos" en el menú lateral, el sistema redirige al módulo de empleos sin errores.']
    ])
    fill_plan_table(doc.tables[2], [
        ['PS-001', 'Sistema', 'FUNC-01', 'Validar manejo de error: Contraseña incorrecta', '25/05/2026', 'El sistema muestra mensaje de error y no redirige al Feed.'],
        ['PS-002', 'Sistema', 'FUNC-02', 'Validar manejo de error: Correo inválido en login (ej: 12345, usuario@)', '25/05/2026', 'El sistema muestra: "Correo no válido. Usa el formato: usuario@dominio.com".'],
        ['PS-003', 'Sistema', 'FUNC-03', 'Validar manejo de error: Publicación vacía', '25/05/2026', 'Al intentar publicar con el campo "Comparte un logro..." vacío, el sistema muestra: "El contenido no puede estar vacío" y no publica.']
    ])
    fill_plan_table(doc.tables[3], [
        ['PI-001', 'Integración', 'INT-01', 'Validar integración: Login → Feed', '25/05/2026', 'Tras login exitoso, el Feed muestra el mensaje de bienvenida y el campo "Comparte un logro...".'],
        ['PI-002', 'Integración', 'INT-02', 'Validar integración: Publicación → Feed', '25/05/2026', 'La publicación creada aparece en el Feed con foto de perfil, nombre de usuario y texto.'],
        ['PI-003', 'Integración', 'INT-03', 'Validar integración: Menú lateral → Módulos', '25/05/2026', 'Al hacer clic en "Conexiones", "Empleos" o "Mensajes" en el menú, el sistema redirige correctamente.']
    ])
    fill_plan_table(doc.tables[4], [
        ['UT-001', 'Unitaria', 'UNIT-01', 'Validar campo de publicación (longitud)', '25/05/2026', 'El campo "Comparte un logro..." no acepta más de 500 caracteres.'],
        ['UT-002', 'Unitaria', 'UNIT-02', 'Validar botón "+ Adjuntar imagen"', '25/05/2026', 'Al hacer clic, el sistema abre el selector de archivos.'],
        ['UT-003', 'Unitaria', 'UNIT-03', 'Validar botón "Cerrar sesión"', '25/05/2026', 'Al hacer clic, el sistema cierra la sesión y redirige al login.']
    ])

    # Move the plan tables after paragraph 120 in the correct order.
    anchor = doc.paragraphs[120]
    for table in [doc.tables[4], doc.tables[3], doc.tables[2], doc.tables[1]]:
        move_table_after(table, anchor)

    # Remove the old redundant table from the original plan section.
    delete_table(doc.tables[5])

    # Update the results section intro.
    set_paragraph(doc.paragraphs[150], '6. RESULTADOS DE LAS PRUEBAS', bold=True)
    set_paragraph(doc.paragraphs[151], 'A continuación, se detallan los resultados de cada prueba ejecutada, organizados por módulo. Cada tabla incluye el resultado, evidencia y observaciones para cada caso de prueba.')

    # Hide the old narrative paragraphs below section 6 so the new evidence tables stand out.
    for idx in range(153, 275):
        doc.paragraphs[idx].text = ''

    # Fill the 12 visible result tables that already exist in the document.
    result_tables = [
        ('PU-001', 'Aprobada', 'PU-001_LoginFeed.png', 'Se verificó que el usuario inicia sesión correctamente y visualiza el Feed con el mensaje de bienvenida y el compositor de publicaciones.', 'Figura 8.1. Acceso correcto al feed principal.'),
        ('PU-002', 'Aprobada', 'PU-002_PublicacionFeed.png', 'La publicación aparece en el Feed con foto de perfil, nombre de usuario y texto visible.', 'Figura 8.2. Publicación visible en el feed.'),
        ('PU-003', 'Aprobada', 'PU-003_NavegacionModulos.png', 'La navegación lateral redirige correctamente al módulo de empleos sin errores visuales.', 'Figura 8.3. Navegación entre módulos.'),
        ('PS-001', 'Aprobada', 'PS-001_ContraseñaIncorrecta.png', 'El sistema muestra un mensaje de error y evita el acceso al Feed.', 'Figura 8.4. Contraseña incorrecta.'),
        ('PS-002', 'Aprobada', 'PS-002_CorreoInvalido.png', 'Se valida el formato del correo y se muestra el mensaje de ayuda esperado.', 'Figura 8.5. Correo inválido en login.'),
        ('PS-003', 'Aprobada', 'PS-003_PublicacionVacia.png', 'Al intentar publicar sin contenido se muestra el mensaje "El contenido no puede estar vacío" y no se publica.', 'Figura 8.6. Publicación vacía.'),
        ('PI-001', 'Aprobada', 'PI-001_LoginFeedIntegracion.png', 'El login exitoso carga el Feed con el mensaje de bienvenida y el campo de publicación.', 'Figura 8.7. Login → Feed.'),
        ('PI-002', 'Aprobada', 'PI-002_PublicacionFeedIntegracion.png', 'La publicación creada se refleja inmediatamente en el Feed con sus datos completos.', 'Figura 8.8. Publicación → Feed.'),
        ('PI-003', 'Aprobada', 'PI-003_MenuModulos.png', 'La navegación lateral abre correctamente el módulo de empleos.', 'Figura 8.9. Menú lateral → módulos.'),
        ('UT-001', 'Aprobada', 'UT-001_CampoPublicacion.png', 'El campo de publicación queda limitado a 500 caracteres y muestra el contador correspondiente.', 'Figura 8.10. Límite de 500 caracteres.'),
        ('UT-002', 'Aprobada', 'UT-002_AdjuntarImagen.png', 'El botón de adjuntar imagen se encuentra visible y operativo dentro del compositor de publicaciones.', 'Figura 8.11. Botón adjuntar imagen.'),
        ('UT-003', 'Aprobada', 'UT-003_CerrarSesion.png', 'Al cerrar sesión, el sistema redirige correctamente al login.', 'Figura 8.12. Cierre de sesión.')
    ]

    for table, (test_id, result, image_name, observations, caption) in zip(doc.tables[6:18], result_tables):
        headers = ['ID de Prueba', 'Resultado', 'Evidencia', 'Observaciones']
        for col, header in enumerate(headers):
            set_table_cell_text(table.rows[0].cells[col], header, font_size=9, bold=True)
        set_table_cell_text(table.rows[1].cells[0], test_id)
        set_table_cell_text(table.rows[1].cells[1], result)
        set_table_cell_text(table.rows[1].cells[3], observations)
        add_image_to_cell(table.rows[1].cells[2], caption, image_name, width=2.25)

    # Remove the extra old result tables that are no longer part of the report.
    for table in list(doc.tables[18:31]):
        delete_table(table)

    doc.save(str(OUTPUT_PATH))
    print(f'Documento actualizado: {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
