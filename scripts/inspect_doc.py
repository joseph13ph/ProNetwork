from docx import Document
from pathlib import Path

p = Path(r'c:\Users\AQuilezZ\OneDrive - Universidad Simón Bolívar\Documentos\7 semestre\TECNOLOGÍAS WEB\corte 3\web\proconnect\INFORME DE PRUEBAS DE SOFTWARE.docx')
doc = Document(str(p))
print('PARAS', len(doc.paragraphs))
print('TABLES', len(doc.tables))
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if t.startswith('6.') or 'RESULTADOS DE LAS PRUEBAS' in t or t.startswith('Módulo '):
        print(f'P{i}: {t}')
for ti, table in enumerate(doc.tables[:8]):
    print('TABLE', ti, 'rows', len(table.rows), 'cols', len(table.columns))
    for ri, row in enumerate(table.rows[:4]):
        vals = [c.text.strip().replace('\n', ' | ')[:180] for c in row.cells]
        print(' ', ri, vals)
