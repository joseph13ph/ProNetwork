from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / 'INFORME DE PRUEBAS DE SOFTWARE.docx'
IMG_DIR = ROOT / 'docs' / 'screenshots'

SECTION4_START = '4. PLAN DE PRUEBAS'
SECTION5_START = '5. EJECUCIÓN DE PRUEBAS'
SECTION6_START = '6. RESULTADOS DE LAS PRUEBAS'
SECTION7_START = '7. RESULTADOS GENERALES DE LA EJECUCIÓN'


def paragraph_text(element, doc):
    if element.tag != qn('w:p'):
        return ''
    return Paragraph(element, doc).text.strip()


def find_block_index(blocks, doc, heading_text):
    for index, block in enumerate(blocks):
        if paragraph_text(block, doc) == heading_text:
            return index
    raise ValueError(f'No se encontró la sección: {heading_text}')


def copy_blocks(blocks):
    return [deepcopy(block) for block in blocks]


def clear_body(body):
    for child in list(body):
        body.remove(child)


def add_paragraph(doc, text, style=None, bold=False, alignment=None):
    paragraph = doc.add_paragraph(text)
    if style:
        paragraph.style = style
    if bold:
        paragraph.runs[0].bold = True
    if alignment is not None:
        paragraph.alignment = alignment
    return paragraph


def set_cell_caption_and_image(cell, caption, image_name, image_width=2.55):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption + '\n')
    run.bold = True
    run.font.size = Pt(8)
    image_path = IMG_DIR / image_name
    if image_path.exists():
        image_run = paragraph.add_run()
        image_run.add_picture(str(image_path), width=Inches(image_width))


def add_table_from_rows(doc, headers, rows, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return table


def add_module_result_table(doc, test_id, result, evidence_caption, image_name, observations):
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    headers = ['ID de Prueba', 'Resultado', 'Evidencia', 'Observaciones']
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    table.rows[1].cells[0].text = test_id
    table.rows[1].cells[1].text = result
    table.rows[1].cells[3].text = observations
    set_cell_caption_and_image(table.rows[1].cells[2], evidence_caption, image_name)
    return table


def build_middle_content():
    temp = Document()

    add_paragraph(temp, SECTION4_START, style='Heading 1')
    add_paragraph(temp, 'Basado en la interfaz real de ProConnect: Feed, Conexiones, Empleos, Mensajes, botón "Publicar ahora" y campo de texto "Comparte un logro, una vacante o una idea que pueda inspirar a otros...".')
    add_paragraph(temp, 'Pruebas de Aceptación (Usuario) - 3 pruebas', style='Heading 2')
    add_paragraph(temp, 'Objetivo: validar flujos completos desde la perspectiva del usuario final.')
    add_table_from_rows(
        temp,
        ['ID de caso de prueba', 'Tipo de prueba', 'ID de paso', 'Descripción del paso', 'Fecha de la prueba', 'Resultado Esperado'],
        [
            ['PU-001', 'Usuario', 'ACC-01', 'Validar flujo completo: Inicio de sesión → Feed', '25/05/2026', 'El sistema redirige al Feed y muestra el mensaje "Bienvenido, [nombre]" con el campo de publicación.'],
            ['PU-002', 'Usuario', 'ACC-02', 'Validar flujo completo: Creación de publicación', '25/05/2026', 'Al hacer clic en "Publicar ahora", la publicación aparece en el feed con foto de perfil, nombre y texto.'],
            ['PU-003', 'Usuario', 'ACC-03', 'Validar flujo completo: Navegación entre módulos', '25/05/2026', 'Al hacer clic en "Empleos" en el menú lateral, el sistema redirige al módulo de empleos sin errores.']
        ]
    )

    add_paragraph(temp, 'Pruebas Funcionales (Sistema) - 3 pruebas', style='Heading 2')
    add_paragraph(temp, 'Objetivo: validar que el sistema rechace acciones inválidas y muestre errores claros.')
    add_table_from_rows(
        temp,
        ['ID de caso de prueba', 'Tipo de prueba', 'ID de paso', 'Descripción del paso', 'Fecha de la prueba', 'Resultado Esperado'],
        [
            ['PS-001', 'Sistema', 'FUNC-01', 'Validar manejo de error: Contraseña incorrecta', '25/05/2026', 'El sistema muestra mensaje de error y no redirige al Feed.'],
            ['PS-002', 'Sistema', 'FUNC-02', 'Validar manejo de error: Correo inválido en login (ej: 12345, usuario@)', '25/05/2026', 'El sistema muestra: "Correo no válido. Usa el formato: usuario@dominio.com".'],
            ['PS-003', 'Sistema', 'FUNC-03', 'Validar manejo de error: Publicación vacía', '25/05/2026', 'Al intentar publicar con el campo "Comparte un logro..." vacío, el sistema muestra: "El contenido no puede estar vacío" y no publica.']
        ]
    )

    add_paragraph(temp, 'Pruebas de Integración - 3 pruebas', style='Heading 2')
    add_paragraph(temp, 'Objetivo: validar la interacción entre componentes de la interfaz.')
    add_table_from_rows(
        temp,
        ['ID de caso de prueba', 'Tipo de prueba', 'ID de paso', 'Descripción del paso', 'Fecha de la prueba', 'Resultado Esperado'],
        [
            ['PI-001', 'Integración', 'INT-01', 'Validar integración: Login → Feed', '25/05/2026', 'Tras login exitoso, el Feed muestra el mensaje de bienvenida y el campo "Comparte un logro...".'],
            ['PI-002', 'Integración', 'INT-02', 'Validar integración: Publicación → Feed', '25/05/2026', 'La publicación creada aparece en el Feed con foto de perfil, nombre de usuario y texto.'],
            ['PI-003', 'Integración', 'INT-03', 'Validar integración: Menú lateral → Módulos', '25/05/2026', 'Al hacer clic en "Conexiones", "Empleos" o "Mensajes" en el menú, el sistema redirige correctamente.']
        ]
    )

    add_paragraph(temp, 'Pruebas Unitarias - 3 pruebas', style='Heading 2')
    add_paragraph(temp, 'Objetivo: validar componentes individuales de la UI.')
    add_table_from_rows(
        temp,
        ['ID de caso de prueba', 'Tipo de prueba', 'ID de paso', 'Descripción del paso', 'Fecha de la prueba', 'Resultado Esperado'],
        [
            ['UT-001', 'Unitaria', 'UNIT-01', 'Validar campo de publicación (longitud)', '25/05/2026', 'El campo "Comparte un logro..." no acepta más de 500 caracteres.'],
            ['UT-002', 'Unitaria', 'UNIT-02', 'Validar botón "+ Adjuntar imagen"', '25/05/2026', 'Al hacer clic, el sistema abre el selector de archivos.'],
            ['UT-003', 'Unitaria', 'UNIT-03', 'Validar botón "Cerrar sesión"', '25/05/2026', 'Al hacer clic, el sistema cierra la sesión y redirige al login.']
        ]
    )

    add_paragraph(temp, SECTION5_START, style='Heading 1')
    add_paragraph(temp, 'La ejecución de pruebas del proyecto ProConnect se realizó en un entorno local para validar el comportamiento de la interfaz real, la navegación entre módulos y las acciones de autenticación, publicación y cierre de sesión.')
    add_paragraph(temp, 'Entorno de ejecución:')
    add_paragraph(temp, 'Frontend: React 18 + Vite, ejecutado en http://localhost:5173.')
    add_paragraph(temp, 'Backend: Node.js + Express, ejecutado en http://localhost:5001.')
    add_paragraph(temp, 'Base de datos: SQLite local en backend/data/proconnect.sqlite.')
    add_paragraph(temp, 'Navegador: Google Chrome.')
    add_paragraph(temp, 'Cobertura: login, feed, navegación entre módulos, publicación de contenido, validaciones y cierre de sesión.')

    add_paragraph(temp, SECTION6_START, style='Heading 1')
    add_paragraph(temp, 'A continuación, se detallan los resultados de cada prueba ejecutada. Cada tabla incluye el resultado, la evidencia y las observaciones del caso de prueba.')

    add_paragraph(temp, 'Pruebas de Aceptación (Usuario)', style='Heading 2')
    add_paragraph(temp, 'Prueba PU-001 – Flujo Completo: Login → Feed', style='Heading 3')
    add_module_result_table(
        temp,
        'PU-001',
        'Aprobada',
        'PU-001_LoginFeed.png',
        'PU-001_LoginFeed.png',
        'Se verificó que el usuario inicia sesión correctamente y visualiza el Feed con el mensaje de bienvenida y el compositor de publicaciones.'
    )
    add_paragraph(temp, 'Prueba PU-002 – Flujo Completo: Creación de Publicación', style='Heading 3')
    add_module_result_table(
        temp,
        'PU-002',
        'Aprobada',
        'PU-002_PublicacionFeed.png',
        'PU-002_PublicacionFeed.png',
        'La publicación aparece en el Feed con foto de perfil, nombre de usuario y texto visible.'
    )
    add_paragraph(temp, 'Prueba PU-003 – Flujo Completo: Navegación entre Módulos', style='Heading 3')
    add_module_result_table(
        temp,
        'PU-003',
        'Aprobada',
        'PU-003_NavegacionModulos.png',
        'PU-003_NavegacionModulos.png',
        'La navegación lateral redirige correctamente al módulo de empleos sin errores visuales.'
    )

    add_paragraph(temp, 'Pruebas Funcionales (Sistema)', style='Heading 2')
    add_paragraph(temp, 'Prueba PS-001 – Manejo de Error: Contraseña Incorrecta', style='Heading 3')
    add_module_result_table(
        temp,
        'PS-001',
        'Aprobada',
        'PS-001_ContraseñaIncorrecta.png',
        'PS-001_ContraseñaIncorrecta.png',
        'El sistema muestra un mensaje de error y evita el acceso al Feed.'
    )
    add_paragraph(temp, 'Prueba PS-002 – Manejo de Error: Correo Inválido en Login', style='Heading 3')
    add_module_result_table(
        temp,
        'PS-002',
        'Aprobada',
        'PS-002_CorreoInvalido.png',
        'PS-002_CorreoInvalido.png',
        'Se valida el formato del correo y se muestra el mensaje de ayuda esperado.'
    )
    add_paragraph(temp, 'Prueba PS-003 – Manejo de Error: Publicación Vacía', style='Heading 3')
    add_module_result_table(
        temp,
        'PS-003',
        'Aprobada',
        'PS-003_PublicacionVacia.png',
        'PS-003_PublicacionVacia.png',
        'Al intentar publicar sin contenido se muestra el mensaje "El contenido no puede estar vacío" y no se publica.'
    )

    add_paragraph(temp, 'Pruebas de Integración', style='Heading 2')
    add_paragraph(temp, 'Prueba PI-001 – Integración: Login → Feed', style='Heading 3')
    add_module_result_table(
        temp,
        'PI-001',
        'Aprobada',
        'PI-001_LoginFeedIntegracion.png',
        'PI-001_LoginFeedIntegracion.png',
        'El login exitoso carga el Feed con el mensaje de bienvenida y el campo de publicación.'
    )
    add_paragraph(temp, 'Prueba PI-002 – Integración: Publicación → Feed', style='Heading 3')
    add_module_result_table(
        temp,
        'PI-002',
        'Aprobada',
        'PI-002_PublicacionFeedIntegracion.png',
        'PI-002_PublicacionFeedIntegracion.png',
        'La publicación creada se refleja inmediatamente en el Feed con sus datos completos.'
    )
    add_paragraph(temp, 'Prueba PI-003 – Integración: Menú Lateral → Módulos', style='Heading 3')
    add_module_result_table(
        temp,
        'PI-003',
        'Aprobada',
        'PI-003_MenuModulos.png',
        'PI-003_MenuModulos.png',
        'La navegación lateral abre correctamente el módulo de empleos.'
    )

    add_paragraph(temp, 'Pruebas Unitarias', style='Heading 2')
    add_paragraph(temp, 'Prueba UT-001 – Validación de Campo de Publicación', style='Heading 3')
    add_module_result_table(
        temp,
        'UT-001',
        'Aprobada',
        'UT-001_CampoPublicacion.png',
        'UT-001_CampoPublicacion.png',
        'El campo de publicación queda limitado a 500 caracteres y muestra el contador correspondiente.'
    )
    add_paragraph(temp, 'Prueba UT-002 – Validación de Botón "+ Adjuntar Imagen"', style='Heading 3')
    add_module_result_table(
        temp,
        'UT-002',
        'Aprobada',
        'UT-002_AdjuntarImagen.png',
        'UT-002_AdjuntarImagen.png',
        'El botón de adjuntar imagen se encuentra visible y operativo dentro del compositor de publicaciones.'
    )
    add_paragraph(temp, 'Prueba UT-003 – Validación de Botón "Cerrar Sesión"', style='Heading 3')
    add_module_result_table(
        temp,
        'UT-003',
        'Aprobada',
        'UT-003_CerrarSesion.png',
        'UT-003_CerrarSesion.png',
        'Al cerrar sesión, el sistema redirige correctamente al login.'
    )

    blocks = []
    for child in temp.element.body.iterchildren():
        if child.tag == qn('w:sectPr'):
            continue
        if child.tag == qn('w:p') and not paragraph_text(child, temp):
            continue
        blocks.append(deepcopy(child))
    return blocks


def main():
    doc = Document(str(DOCX_PATH))
    body = doc.element.body
    original_children = list(body.iterchildren())
    sect_pr = None
    if original_children and original_children[-1].tag == qn('w:sectPr'):
        sect_pr = deepcopy(original_children[-1])

    section4_idx = find_block_index(original_children, doc, SECTION4_START)
    section5_idx = find_block_index(original_children, doc, SECTION5_START)
    section7_idx = find_block_index(original_children, doc, SECTION7_START)

    before_section4 = copy_blocks(original_children[:section4_idx])
    after_section7 = copy_blocks(original_children[section7_idx:])
    middle_blocks = build_middle_content()

    clear_body(body)
    for block in before_section4:
        body.append(block)
    for block in middle_blocks:
        body.append(block)
    for block in after_section7:
        body.append(block)
    if sect_pr is not None:
        body.append(sect_pr)

    doc.save(str(DOCX_PATH))
    print(f'Documento actualizado: {DOCX_PATH}')


if __name__ == '__main__':
    main()
